# 普源精电科技技术有限公司版权所有 (2021 - )

# !/usr/bin/env python3
# _*_ coding: utf-8 _*_
"""
源文件名: open_source.py
功能描述: 
作者: ShuJie
版本: Ver2.0.1
创建日期: 2021/9/29 13:39

修改历史:
修改日期：
版本号：
修改人：
修改内容：
"""

# 导入的包
import os
import chardet
import time
from tqdm import tqdm
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # qn设置中文字体
from docx.oxml.ns import nsdecls  # 背景色使用
from docx.oxml import parse_xml   # 背景色使用
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT   # 字体对齐
from docx.enum.table import WD_TABLE_ALIGNMENT

import color_list as cl  # 导入颜色进制值,位于color_list.py

# 宏定义
USE_DEBUG = 1      # 调试信息

# 版本号
VERSION = '2.0.1'

# 函数实现


"""
主函数名: tableMain
功能描述: 生成copyright和license表格 
         对已存在的文件插入表格操作, 如新建word则不用此方案
         表格不可以为空, 否则会报错
         对表格大小最好确定rows和cols再调用
作者: ShuJie
版本: Ver1.0
创建日期: 2021/9/29 13:39

修改历史:
修改日期：
版本号：
修改人：
修改内容：
"""

# 设置表格的边框---->由于现有API无法对没有表格的docx进行属性设置, 所以采用自添加
def setCellBorder(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    setCellBorder(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existence, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existence, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

# 设置表格某行背景色
def tabBgColor(table, rows, cols, Color):
    """

    :param table:  table对象
    :param rows:   行数
    :param cols:   列数
    :param Color:  背景色
    :return:       无
    """

    shading_list = locals()
    color = cl.colorNamesDict[Color] if Color is not None and cl.colorNamesDict[Color] is not None \
        else cl.colorNamesDict['white']    # 不指定默认白色
    for i in range(cols):
        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'
                                                        .format(nsdecls('w'), bgColor=Color))
        table.rows[rows].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])

# 添加表格
def addTable(docObj, rows, cols, Color=None):
    """

    :param docObj: word文件对象
    :param rows:   行数
    :param cols:   列数
    :param Color:  表格边框颜色
    :return:       table对象
    """

    table = docObj.add_table(rows, cols)
    color = cl.colorNamesDict[Color] if Color is not None and cl.colorNamesDict[Color] is not None \
        else cl.colorNamesDict['black']       # 不指定颜色默认为黑色
    for i in range(rows):
        rowCell = table.rows[i].cells
        for j in range(cols):
            setCellBorder(rowCell[j],
                          top={"sz": 12, "val": "single", "color": color},
                          bottom={"sz": 12, "val": "single", "color": color},
                          left={"sz": 12, "val": "single", "color": color},
                          right={"sz": 12, "val": "single", "color": color},
                          )
    return table

# 添加license表格,请保证至少3列
def addLicenseTable(docObj, rows, cols, textList=None):
    """

    :param docObj:   word文件对象
    :param rows:     行数
    :param cols:     列数
    :param textList: 插入内容,格式为[[package, version, license], ...]...
    :return:         无
    """

    p4 = docObj.add_paragraph('表 license')
    p4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER   # 貌似无效
    table = addTable(docObj, rows, cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER                     # 貌似无效
    tabBgColor(table, 0, 3, 'gray')        # 添加背景色
    table.cell(0, 0).text = 'Package'      # 根据开源声明表格设定
    table.cell(0, 1).text = 'Version'
    table.cell(0, 2).text = 'License'
    if textList is None:
        textList = []
    # 如有内容, 插入表格
    for i in range(len(textList)):
        for j in range(len(textList[i])):
            table.cell(i + 1, j).text = textList[i][j]

# 添加copyright表格,请保证至少3列
def addCopyrightTable(docObj, rows, cols, textList=None):
    """

    :param docObj:   word文件对象
    :param rows:     行数
    :param cols:     列数
    :param textList: 插入内容,格式为 格式为[[package, version, copyright], ...]...
    :return:         无
    """

    p4 = docObj.add_paragraph('表 copyright')
    p4.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table = addTable(docObj, rows, cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabBgColor(table, 0, 3, 'gray')        # 添加背景色
    table.cell(0, 0).text = 'Package'      # 根据开源声明表格设定
    table.cell(0, 1).text = 'Version'
    table.cell(0, 2).text = 'Copyright'
    if textList is None:
        textList = []
    # 如有内容, 插入表格
    for i in range(len(textList)):
        for j in range(len(textList[i])):
            table.cell(i + 1, j).text = textList[i][j]

# 表格生成主函数
def tableMain(fileName, copyList=None, licenseList=None):
    """

    :param fileName:    已存在的文件名
    :param copyList:    copyright列表,格式为[[package, version, copyright], ...]...
    :param licenseList: license列表,格式为[[package, version, license], ...]...
    :return:            无
    """

    file = Document(fileName)             # 文件对象
    if copyList is None:                  # copyright列表为空时
        copyList = []
    if licenseList is None:               # license列表为空时
        licenseList = []
    if len(copyList) > 0:                 # 有数据, 创建copyright表格
        addCopyrightTable(file, len(copyList) + 1, 3, textList=copyList)
        if USE_DEBUG:
            print('create table successfully!')
    else:
        print('no copyright data. Can\'t create a table.')
    if len(licenseList) > 0:              # 有数据, 创建license表格
        addLicenseTable(file, len(licenseList) + 1, 3, textList=licenseList)
        if USE_DEBUG:
            print('create table successfully!')
    else:
        print('no license data. Can\'t create a table.')
    file.save(fileName)


"""
主函数名: insertMain
功能描述: 为文件添加copyright和license
作者: ShuJie
版本: Ver1.0
创建日期: 2021/9/30 10:17

修改历史:   修改次数一
修改日期：  2021年10月12日
版本号：    Ver2.0
修改人：    ShuJie
修改内容：   增加所有文件查询函数
"""

# 判断文件编码
def checkFileEncode(src):
    with open(src, 'rb') as f:
        return chardet.detect(f.read()).get('encoding')

# 查找指定类型文件
def findFileInSuffix(suffix, rootDir=None, ignore=None):
    """

    :param rootDir: 指定目录
    :param suffix:  指定类型, 例如txt, c, cpp
    :param ignore:  忽略的目录
    :return:        结果路径list
    """

    if rootDir is None:        # 路径为空, 默认当前路径
        rootDir = './'
        rootName = 'curr'      # 当前路径
    else:
        rootName = rootDir.split('\\')[-1]
    if ignore is None:
        ignore = ['.git', 'venv', '.repo', '.idea']    # 默认忽略git, repo
    else:
        ignore.append('.git')
        ignore.append('.repo')
    if USE_DEBUG:
        print('当前搜索路径:', rootDir)
    result = []

    txtName = rootName + '_file_in_suffix.txt'      # 本函数输出记录文件
    if os.path.exists(txtName):
        if time.time() - os.stat(txtName).st_mtime > 30 * 60:              # 修改时间大于30分钟才会去覆盖文件
            # 修改时间大于30分钟, 覆盖文件
            if USE_DEBUG:
                print('覆盖文件:', txtName)
            with open(txtName, 'w') as f:
                for name in suffix:
                    suffixName = "." + name
                    walk = os.walk(rootDir)
                    for root, dirs, files in walk:
                        if len(files) < 1:
                            continue
                        for i in ignore:
                            if i in dirs:
                                dirs.remove(i)
                        for file in files:
                            if USE_DEBUG:
                                print('reading', os.path.join(root, file))
                            fileName, serSuffixName = os.path.splitext(file)
                            if serSuffixName == suffixName:
                                result.append(os.path.join(root, file))
                                f.write(os.path.join(root, file) + '\n')
        else:
            # 修改时间不足, 直接读文件
            if USE_DEBUG:
                print('读取文件:', txtName)
            with open(txtName, 'r') as f:
                line = f.readlines()
                for i in line:
                    result.append(i.strip('\n'))
    # 输出文件不存在, 创建
    else:
        with open(txtName, 'w') as f:
            for name in suffix:
                suffixName = "." + name
                walk = os.walk(rootDir)
                for root, dirs, files in walk:
                    if len(files) < 1:
                        continue
                    for i in ignore:
                        if i in dirs:
                            dirs.remove(i)
                    for file in files:
                        if USE_DEBUG:
                            print('reading', os.path.join(root, file))
                        fileName, serSuffixName = os.path.splitext(file)
                        if serSuffixName == suffixName:
                            result.append(os.path.join(root, file))
                            f.write(os.path.join(root, file) + '\n')
    return result

# 查找所有类型文件
def findAllFile(rootDir=None, ignore=None):
    """

    :param rootDir:   查询路径,为空默认为当前路径
    :param ignore:    忽略文件或者目录
    :return:          结果路径list
    """

    if rootDir is None:        # 路径为空, 默认当前路径
        rootDir = './'
        rootName = 'curr'      # 当前路径
    else:
        rootName = rootDir.split('\\')[-1]
    if ignore is None:
        ignore = ['.git', 'venv', '.repo', '.idea']    # 默认忽略git, repo
    else:
        ignore.append('.git')
        ignore.append('.repo')
    if USE_DEBUG:
        print('当前搜索路径:', rootDir)
    result = []

    txtName = rootName + 'all_file.txt'
    if os.path.exists(txtName):
        if time.time() - os.stat(txtName).st_mtime > 30 * 60:  # 修改时间大于30分钟才会去覆盖文件
            # 修改时间大于30分钟, 覆盖文件
            if USE_DEBUG:
                print('覆盖文件:', txtName)
            with open(txtName, 'w') as f:
                walk = os.walk(rootDir)
                for root, dirs, files in walk:
                    if len(files) < 1:
                        continue
                    for i in ignore:
                        if i in dirs:
                            dirs.remove(i)
                    for file in files:
                        if USE_DEBUG:
                            print('reading', os.path.join(root, file))
                        result.append(os.path.join(root, file))
                        f.write(os.path.join(root, file) + '\n')
        else:
            # 修改时间不足, 直接读文件
            if USE_DEBUG:
                print('读取文件:', txtName)
            with open(txtName, 'r') as f:
                line = f.readlines()
                for i in line:
                    result.append(i.strip('\n'))
        # 输出文件不存在, 创建
    else:
        with open(txtName, 'w') as f:
            walk = os.walk(rootDir)
            for root, dirs, files in walk:
                if len(files) < 1:
                    continue
                for i in ignore:
                    if i in dirs:
                        dirs.remove(i)
                for file in files:
                    if USE_DEBUG:
                        print('reading', os.path.join(root, file))
                    result.append(os.path.join(root, file))
                    f.write(os.path.join(root, file) + '\n')
    return result

# 查找指定字符串
def findFileInString(fileList, string):

    """

    :param fileList:  查找的文件目录
    :param string:    指定的字符串list
    :return:          路径及字符串内容list
    """
    result = []
    for i in tqdm(range(len(fileList))):
        count = 0                                  # 计数器
        code = checkFileEncode(fileList[i])
        with open(fileList[i], 'r', encoding=code, errors='ignore') as f:    # 忽略中文带来的影响
            for line in f.readlines():
                for j in string:
                    if j in line:
                        data = fileList[i] + ':' + line
                        if fileList[i] + ':' + line not in result:         # 避免重复
                            result.append(data)
                            count += 1
            if count < len(string):                  # 如果有指定字符串,则count至少为fileList长度
                result.append(fileList[i] + ':' + 'None')

    return result

# copyright 和 license添加
def checkCopyrightAndLicense(srcList, compareList=None):
    # 备注: 根据未扫描到相应字符串的文件进行copyright和license分析
    """

    :param srcList:        源目录list
    :param compareList:    对比的目录list,用于copyright和license查询时索引已查找过的记录
    :return:               无
    """
    scanOfString = ['Copyright', 'License']
    rigolTitle = '普源精电'           # 普源公司专属文件编写
    if compareList is None:
        compareFlag = 0     # 对比标志位,若为0,则表明没有用于查询的列表
    else:
        compareFlag = 1
    scanOfFile = findFileInString(srcList, scanOfString)   # 扫描结果
    for string in scanOfFile:
        path = string.split(':')[0]      # 默认使用相对路径, 从而拆分路径,若分在不同的磁盘,则会出现问题
        comment = string.split(':')[-1]  # 文件扫描内容
        fileName = path.split('\\')[-1]  # 文件名
        rigolFileFlag = 0                # 普源公司文件所有标志位
        if comment == 'None':            # None说明未有copyright和license
            code = checkFileEncode(path)
            with open(path, 'r', encoding=code) as f:
                for line in f.readlines():
                    if rigolTitle in line:     # 普源公司员工编写
                        rigolFileFlag = 1
                        break
            with open(path, 'r+', encoding=code) as f:
                if rigolFileFlag:             # 普源公司copyright
                    old = f.read()
                    f.seek(0)
                    rigolCopyright = '/*\n' \
                                     f' * {fileName}\n' \
                                     ' *\n * Copyright (c) 2018 RIGOL,inc\n *\n' \
                                     ' * This program is free software; you can redistribute it and/or modify it\n' \
                                     ' * under the terns and conditions of the GNU General Public License,\n' \
                                     ' * version 2, as published by the Free Software Foundation.\n */\n\n'
                    f.write(rigolCopyright)
                    f.write(old)
                else:
                    pass

# 添加copyright和license主函数
def insertMain(path=None, mode=None):
    """

    :param mode:  查询模式,为空默认查询所有,为specific为指定后缀名的文件
    :param path:  查询路径,为空默认为当前路径
    :return:      无
    """

    searchSuffix = ['c', 'h', 'cpp', 'py', 'java']            # 查找的源文件类型
    if mode is None:
        searchList = findAllFile(rootDir=path)
    else:
        searchList = findFileInSuffix(searchSuffix, path)
    checkCopyrightAndLicense(searchList)


if __name__ == '__main__':
    copList = [['linux', '3.2.0', 'GPL 2.0'], ['NE10', '1.1.1', 'GPL 2.0']]
    tableMain('a.docx', copyList=copList)
    insertMain(path='.\\key_check')
