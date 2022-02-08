# 普源精电科技技术有限公司版权所有 (2021 - )

# !/usr/bin/env python3
# _*_ coding: utf-8 _*_
"""
源文件名: open_source.py
功能描述: 
作者: ShuJie
版本: Ver1.0
创建日期: 2021/9/28 11:32

修改历史:
修改日期：
版本号：
修改人：
修改内容：
"""

import docx
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from os.path import abspath
from win32com import client


class WordFormatInsert:
    """
    该类使用时
    """

    __ERROR_TYPE = 0  # 错误类型 0表示无错误
    __PICTURE_INSERT_ERROR = 0  # 图片插入错误

    def __init__(self, file=None, picture=None):
        self.picture = picture
        if file is not None:
            # 创建文件对象
            try:
                self.doc = Document(file)
            except docx.opc.exceptions.PackageNotFoundError:
                print('file is not existed!')
                print('create new file instead!')
                self.doc = Document()
        else:
            del self

    def __del__(self):
        print('delete object!')

    # 插入图片 默认尺寸200
    def picInsert(self, width=200):
        if self.picture is not None:
            self.doc.add_picture(self.picture, width=Pt(width))
        else:
            print('there is not picture to insert!')


# a = WordFormatInsert(file='aa.docx', picture='1.jpg')
# a.picInsert(width=300)
# document = Document()
# document = Document('aa.docx')
# b = document.add_table(cols=2, rows=2, style='Medium Shading 1 Accent 1')
# b.cell(1, 1).text = 'aaa'
# document.save('aa.docx')

def main(files, final_docx):
    # 启动word应用程序
    word = client.gencache.EnsureDispatch("Word.Application")
    word.Visible = True
    # 新建空白文档
    new_document = word.Documents.Add()
    for fn in files:
        # 打开要合并的每个文件，复制其中的内容到剪切板，然后关闭文件
        fn = abspath(fn)
        temp_document = word.Documents.Open(fn)
        word.Selection.WholeStory()
        word.Selection.Copy()
        temp_document.Close()
        # 粘贴到新文档的最后
        new_document.Range()
        word.Selection.Delete()
        word.Selection.Paste()
    # 保存最终文件，关闭Word应用程序
    new_document.SaveAs(final_docx)
    new_document.Close()
    word.Quit()


List = ['a.docx', 'aa.docx']
main(List, 'out.docx')